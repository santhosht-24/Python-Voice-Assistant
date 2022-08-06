import speech_recognition as sr
import pyttsx3
import datetime
import webbrowser
import wikipedia
import pyjokes
import docx
import os
from openpyxl import load_workbook

r = sr.Recognizer()
engine = pyttsx3.init()
engine.setProperty("rate", 150)

def voiceOutput(command):
    engine.say(command)
    engine.runAndWait()

def greetings():

    hour = int(datetime.datetime.now().hour)
    if 0 <= hour < 12:
        print("Good Morning! ")
        voiceOutput("Good Morning!")

    elif 12 <= hour < 18:
        print("Good Afternoon! ")
        voiceOutput("Good Afternoon!")

    else:
        print("Good Evening! ")
        voiceOutput("Good Evening!")

    print("I am your virtual assistant. What can I do for you today?")
    voiceOutput("I am your virtual Assistant. What can I do for you today?")


def greetingHello():
    print("Hello. I am your virtual assistant. How can I assist you today?")
    voiceOutput("Hello! I am your virtual Assistant. How can I assist you today?")
    run_assistant()

def take_command():
    try:
        with sr.Microphone() as source1:
            r.adjust_for_ambient_noise(source1, duration=0.2)

            print('Getting input...')
            voiceOutput("Getting input.")
            audio1 = r.listen(source1)

            print('Grabbing results...')
            voiceOutput("Grabbing results")
            MyText = r.recognize_google(audio1)

            command = MyText.lower()
            print("Your command : " + command)

    except sr.RequestError as e:
        print("Could not request results; {0}".format(e))

    except sr.UnknownValueError:
        print("unknown error occurred")

    return command

def tellDay():
    day = datetime.datetime.today().weekday() + 1

    Day_dict = {1: 'Monday', 2: 'Tuesday', 3: 'Wednesday', 4: 'Thursday', 5: 'Friday', 6: 'Saturday', 7: 'Sunday'}

    if day in Day_dict.keys():
        day_of_the_week = Day_dict[day]
        print("The day is " + day_of_the_week + " today")
        voiceOutput("Today is " + day_of_the_week)

def tellTime():
    time = str(datetime.datetime.now())
    print(time)
    hour = time[11:13]
    minutes = time[14:16]
    voiceOutput("The time is " + hour + "Hours and" + minutes + "Minutes")

def searchWiki(command):
    print("Processing input...")
    voiceOutput("Processing input.")

    command = command.replace("from ", "")
    command = command.replace("wikipedia", "")
    command = command.replace("search ", "")
    command = command.replace("for ", "")

    result = wikipedia.summary(command, sentences=4)
    print("According to wikipedia...")
    voiceOutput("According to wikipedia.")
    print(result)
    voiceOutput(result)


def nameTell(command):
    command = command.replace("my name is ", "")
    print("Hello " + command + ". How are you?")
    voiceOutput("Hello " + command + ". How are you?")

def openClassroom():
    print("Opening Google Classroom...")
    voiceOutput("Opening Google Classroom.")
    webbrowser.open("https://classroom.google.com")

def weather():
    print("Getting weather forecast...")
    voiceOutput("Getting weather forecast.")
    webbrowser.open("https://www.google.com/search?q=weather+forecast")

def googleSearch(command):
    command = command.replace("google", "")
    command = command.replace("search ", "")
    command = command.replace("for ", "")
    command = command.replace("in ", "")
    print("Searching Google for : " + command)
    voiceOutput("Searching Google for " + command)
    webbrowser.open("www.google.com/search?q=" + command)

def whoami():
    print("You are the most wonderful programmer. You are Mr/Mrs Whatsoever. And you are my god because you created me!")
    voiceOutput("You are the most wonderful programmer. You are Mister or Mistress Whatsoever. And you are my god because you created me!")

def tellAJoke():
    joke = pyjokes.get_joke(language="en", category="neutral")
    print("Here's a joke for you...")
    voiceOutput("here is a joke for you.")
    print(joke + " Ha ha ha!")
    voiceOutput(joke + " Ha ha ha!")

def openMyComp():
    print("Opening My Computer...")
    voiceOutput("Opening My Computer.")
    os.startfile("C:/Users/Siva/Desktop/Computer - Shortcut.lnk")

def openCalc():
    print("Opening Calculator...")
    voiceOutput("Opening calculator.")
    os.startfile("C:/Windows/system32/calc.exe")

def openMSword():
    print("Opening Microsoft Office Word...")
    voiceOutput("Opening Microsoft Office Word.")
    os.startfile("C:/ProgramData/Microsoft/Windows/Start Menu/Programs/Microsoft Office/Microsoft Office Word 2007.lnk")

def openMSExcel():
    print("Opening Microsoft Office Excel...")
    voiceOutput("Opening Microsoft Office Excel.")
    os.startfile("C:/ProgramData/Microsoft/Windows/Start Menu/Programs/Microsoft Office/Microsoft Office Excel 2007.lnk")

def openMSppt():
    print("Opening Microsoft Office PowerPoint...")
    voiceOutput("Opening Microsoft Office Power point.")
    os.startfile("C:/ProgramData/Microsoft/Windows/Start Menu/Programs/Microsoft Office/Microsoft Office PowerPoint 2007.lnk")

def workWithExcel():
    print("Name the excel file that you need to work with...")
    voiceOutput("Name the excel file that you need to work with.")

    command = take_command()
    filename = command.lower()

    workbook = load_workbook(filename="C:/Users/Siva/Desktop/" + filename + ".xlsx")
    sheet = workbook.active

    print("Opening " + filename + " worksheet...")
    voiceOutput("Opening " + filename + " worksheet")

    print("Would you like to display the records or edit a record?")
    voiceOutput("Would you like to display the records or edit a record?")

    command = take_command()

    if 'display' in command:

        print("Would you like to display all the records or display rows with particular word?")
        voiceOutput("Would you like to display all the records or display rows with particular word?")

        command = take_command()

        if 'all the records' in command:
            for value in sheet.iter_rows(min_row=1, max_row=11, min_col=1, max_col=4, values_only=True):
                print(value)

        elif 'particular' in command:

            while(1):

                print("Tell me the word to be searched...")
                voiceOutput("Tell me the word to be searched")

                command = take_command()
                searchWord = command
                found = 0
                for value in sheet.iter_rows(min_row=1, max_row=11, min_col=1, max_col=4, values_only=True):
                    if searchWord in (str(value)).lower():
                        print(value)
                        found = 1
                if found == 1:
                    break
                else:
                    continue

    elif 'edit' in command:

        print("Do you wish to Enter new values or edit and replace an existing value?")
        voiceOutput("Do you wish to Enter new values or edit and replace an existing value?")

        command = take_command()

        if 'replace' in command:
            print("Tell the word to be found and replaced...")
            voiceOutput("Tell the word to be found and replaced.")

            command = take_command()
            inputWord = command.lower()

            found = 0
            foundRow = 0
            foundCol = 0
            for row in range(1, sheet.max_row):
                for col in range(1, sheet.max_column):
                    if inputWord == str(sheet.cell(row=row, column=col).value).lower():
                        found = 1
                        foundRow = row
                        foundCol = col
                        break

            if found == 1:
                print("Data found. Tell the word to be replaced in this position")
                voiceOutput("Data found. Tell the word to be replaced in this position.")

                command = take_command()

                sheet.cell(row=foundRow, column=foundCol).value = command

            else:
                print("Data not found!")
                voiceOutput("Data not found")

        elif 'new value' in command:
            print("Give column value to insert data...")
            voiceOutput("Give column value to insert data.")

            command = take_command()

            col = ord(command) - 96

            print("Give row value to insert data...")
            voiceOutput("Give row value to insert data.")

            command = take_command()

            row = int(command)

            print("Tell the data to be inserted in the cell ", chr(col + 96).upper(), row, "...")
            voiceOutput("Tell the data to be inserted in the cell " + chr(col + 96) + "," + str(row))

            command = take_command()

            sheet.cell(row=row, column=col).value = command

        print("Displaying changed Worksheet...")
        voiceOutput("Displaying changed Worksheet.")

        for value in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column, values_only=True):
            print(value)

        workbook.save(filename="C:/Users/Siva/Desktop/" + filename + ".xlsx")

def wordProcess():

    print('Insert text or Find a word in word file...')
    voiceOutput('Insert text or Find a word in word file')

    command = take_command()

    if 'insert' in command:
        doc = docx.Document()
        print('Dictate the paragraph to be entered...')
        voiceOutput('Dictate the paragraph to be entered')
        command = take_command()
        doc.add_paragraph(command)
        print('Data successfully inserted...')
        voiceOutput('Data successfully inserted')
        doc.save('C:/Users/Siva/Desktop/Sample Document.docx')

    if 'find' in command:

        print("Tell the word to be found...")
        voiceOutput("Tell the word to be found")

        command = take_command()
        checkWord = command.lower

        doc = docx.Document('C:/Users/Siva/Desktop/Sample Document.docx')

        paraList = []
        found = 0

        for para in doc.paragraphs:
            paraList.append(para.text)

        for para in paraList:
            sen = para.split(" ")
            for word in sen:
                if word.lower() == checkWord:
                    found = 1
                    print('Word found in the paragraph...')
                    voiceOutput('Word found in the paragraph')
                    print(para)
        if found == 0:
            print('Word not found...')
            voiceOutput('Word not found')

def run_assistant():

    while 1:

        command = take_command()

        if 'what day' in command:
            tellDay()

        elif 'time' in command:
            tellTime()

        elif 'open google' in command:
            print("Opening Google...")
            voiceOutput("Opening Google.")
            webbrowser.open("www.google.com")

        elif 'open wikipedia' in command:
            print("Opening Wikipedia...")
            voiceOutput("Opening Wikipedia.")
            webbrowser.open("www.wikipedia.org")

        elif 'from wikipedia' in command:
            searchWiki(command)

        elif 'my name' in command:
            nameTell(command)

        elif 'open classroom' in command:
            openClassroom()

        elif 'hello' in command:
            greetingHello()

        elif 'weather' in command:
            weather()

        elif 'google search' in command:
            googleSearch(command)

        elif 'search google' in command:
            googleSearch(command)

        elif 'who am i' in command:
            whoami()

        elif 'tell me a joke' in command:
            tellAJoke()

        elif 'open my computer' in command:
            openMyComp()

        elif 'open calculator' in command:
            openCalc()

        elif 'open word' in command:
            openMSword()

        elif 'word' in command:
            wordProcess()

        elif 'open excel' in command:
            openMSExcel()

        elif 'edit excel' in command:
            workWithExcel()

        elif 'open power point' in command:
            openMSppt()

        elif 'news for today' in command:
            print("Getting today's new for you...")
            voiceOutput("Getting today's new for you.")
            webbrowser.open("https://news.google.com")

        elif 'goodbye' in command:
            print("Goodbye. Have a nice day!!")
            voiceOutput("Goodbye. Have a nice day!")
            break


greetings()
run_assistant()



